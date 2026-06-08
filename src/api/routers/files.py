"""文件管理 API 路由"""
from __future__ import annotations

import os
import shutil
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

from fastapi import APIRouter, Form, Header, HTTPException, UploadFile
from fastapi.responses import FileResponse
from pydantic import BaseModel, Field
from starlette.background import BackgroundTask

from config import load_config
from core.storage import (
    build_blob_name,
    delete_file_from_storage,
    download_file_to_local,
    upload_stream_to_storage,
    oss_storage_enabled,
)
from db.auth_context import resolve_user_from_authorization
from utils.desktop_runtime import get_desktop_local_library
from db.session_repository import (
    add_session_file,
    delete_session_file,
    get_session_by_id,
    get_session_files,
    update_file_selection,
)

router = APIRouter(prefix="/api/sessions/{session_id}/files", tags=["文件管理"])
temp_router = APIRouter(prefix="/api/sessions/{session_id}/temp-files", tags=["临时文件管理"])

# 文件存储目录
UPLOAD_DIR = Path("workspace/uploads")
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

# 临时文件存储目录
TEMP_UPLOAD_DIR = Path("workspace/temp_uploads")
TEMP_UPLOAD_DIR.mkdir(parents=True, exist_ok=True)


class SessionFile(BaseModel):
    id: int
    file_name: str
    file_path: str
    file_type: str
    file_size: int
    is_selected: bool
    created_at: str
    storage_key: Optional[str] = None


class FileListResponse(BaseModel):
    data_files: List[SessionFile]
    template_files: List[SessionFile]


class FileSelectionRequest(BaseModel):
    file_id: int
    is_selected: bool


class TempFileResponse(BaseModel):
    """临时文件响应"""
    file_name: str
    file_path: str
    file_type: str
    file_size: int
    storage_key: Optional[str] = None
    created_at: str


def _resolve_current_user(authorization: Optional[str], cfg):
    if not authorization:
        if cfg.auth.require_auth:
            raise HTTPException(status_code=401, detail="需要登录后访问")
        return None
    try:
        return resolve_user_from_authorization(authorization, cfg, required=True)
    except PermissionError as exc:
        raise HTTPException(status_code=401, detail=str(exc))


def _file_to_dict(f) -> Dict[str, Any]:
    return {
        "id": f.id,
        "file_name": f.file_name,
        "file_path": getattr(f, "storage_key", None) or f.file_path,
        "file_type": f.file_type,
        "file_size": f.file_size,
        "is_selected": f.is_selected,
        "created_at": f.created_at.isoformat() + "Z" if f.created_at else "",
        "storage_key": getattr(f, "storage_key", None),
    }


def _cleanup_blob_cache_file(path: Path, cache_root: Path) -> None:
    """仅清理 oss_cache 下的临时下载文件。"""
    try:
        path = path.resolve()
        cache_root = cache_root.resolve()
    except Exception:
        return

    if cache_root not in path.parents:
        return

    try:
        if path.exists() and path.is_file():
            path.unlink()
    except Exception:
        return

    # 尝试向上清理空目录，最多到 cache_root
    current = path.parent
    while current != cache_root and cache_root in current.parents:
        try:
            current.rmdir()
            current = current.parent
        except OSError:
            break


@router.get("", response_model=FileListResponse)
async def list_files(session_id: str, authorization: Optional[str] = Header(default=None)):
    """获取会话的所有文件（按类型分组）"""
    cfg = load_config()
    current_user = _resolve_current_user(authorization, cfg)
    
    # 检查会话是否存在
    session = get_session_by_id(session_id, config=cfg, user_id=current_user.id if current_user else None)
    if not session:
        raise HTTPException(status_code=404, detail="会话不存在")
    
    all_files = get_session_files(session_id, config=cfg, user_id=current_user.id if current_user else None)

    data_files = [_file_to_dict(f) for f in all_files if f.file_type == "data"]
    template_files = [_file_to_dict(f) for f in all_files if f.file_type == "template"]
    
    return FileListResponse(data_files=data_files, template_files=template_files)


@router.post("", response_model=SessionFile)
async def upload_file(
    session_id: str,
    file: UploadFile,
    file_type: str = Form(..., description="文件类型: data 或 template"),
    authorization: Optional[str] = Header(default=None),
):
    """
    上传文件到会话
    """
    cfg = load_config()
    current_user = _resolve_current_user(authorization, cfg)
    
    # 检查会话是否存在
    session = get_session_by_id(session_id, config=cfg, user_id=current_user.id if current_user else None)
    if not session:
        raise HTTPException(status_code=404, detail="会话不存在")
    
    # 验证文件类型
    if file_type not in ("data", "template"):
        raise HTTPException(status_code=400, detail="file_type 必须是 'data' 或 'template'")
    
    # 保存文件
    file_name = file.filename or "unnamed"
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_name = f"{timestamp}_{file_name}"
    file_size = 0
    storage_key = None
    if oss_storage_enabled(cfg):
        storage_key = upload_stream_to_storage(
            file.file,
            config=cfg,
            blob_name=build_blob_name(session_id, safe_name, prefix=cfg.storage.object_key_prefix),
            content_type=file.content_type,
        )
        if not storage_key:
            raise HTTPException(status_code=502, detail="OSS 上传失败")
        try:
            file.file.seek(0, 2)
            file_size = file.file.tell()
            file.file.seek(0)
        except Exception:
            file_size = 0
    else:
        file_path = UPLOAD_DIR / session_id / safe_name
        file_path.parent.mkdir(parents=True, exist_ok=True)
        with open(file_path, "wb") as buffer:
            while chunk := file.file.read(8192):
                buffer.write(chunk)
                file_size += len(chunk)
        storage_key = str(file_path)
    
    # 保存到数据库
    session_file = add_session_file(
        session_id=session_id,
        file_name=file_name,
        file_type=file_type,
        file_path=storage_key or "",
        file_size=file_size,
        config=cfg,
        user_id=current_user.id if current_user else None,
        source="upload",
        role="source",
        storage_key=storage_key,
    )
    
    return _file_to_dict(session_file)


@router.patch("/selection")
async def update_file_selections(
    session_id: str,
    selections: List[FileSelectionRequest],
    authorization: Optional[str] = Header(default=None),
):
    """批量更新文件勾选状态

    请求体格式:
    [
        {"file_id": 1, "is_selected": true},
        {"file_id": 2, "is_selected": false}
    ]
    """
    cfg = load_config()
    current_user = _resolve_current_user(authorization, cfg)

    # 检查会话是否存在
    session = get_session_by_id(session_id, config=cfg, user_id=current_user.id if current_user else None)
    if not session:
        raise HTTPException(status_code=404, detail="会话不存在")

    # 如果传入的是包装对象，尝试解包
    if len(selections) == 1 and hasattr(selections[0], '__dict__'):
        data = selections[0].__dict__
        if 'selections' in data or 'items' in data:
            items = data.get('selections') or data.get('items') or []
            if isinstance(items, list) and len(items) > 0 and isinstance(items[0], dict):
                selections = [FileSelectionRequest(**item) for item in items]

    results = []
    for sel in selections:
        success = update_file_selection(sel.file_id, sel.is_selected, config=cfg, user_id=current_user.id if current_user else None)
        results.append({"file_id": sel.file_id, "success": success})

    return {"results": results}


@router.delete("/{file_id}")
async def delete_file(session_id: str, file_id: int, authorization: Optional[str] = Header(default=None)):
    """删除文件"""
    cfg = load_config()
    current_user = _resolve_current_user(authorization, cfg)
    
    # 检查会话是否存在
    session = get_session_by_id(session_id, config=cfg, user_id=current_user.id if current_user else None)
    if not session:
        raise HTTPException(status_code=404, detail="会话不存在")
    
    # 获取文件信息（用于删除物理文件）
    files = get_session_files(session_id, config=cfg, user_id=current_user.id if current_user else None)
    file_info = next((f for f in files if f.id == file_id), None)
    
    if not file_info:
        raise HTTPException(status_code=404, detail="文件不存在")
    
    # 删除 OSS 对象
    storage_key = getattr(file_info, "storage_key", None)
    if storage_key:
        delete_file_from_storage(storage_key, config=cfg)
    
    # 删除数据库记录
    success = delete_session_file(file_id, config=cfg, user_id=current_user.id if current_user else None)
    
    return {"success": success}


@router.get("/{file_id}/download")
async def download_file(session_id: str, file_id: int, authorization: Optional[str] = Header(default=None)):
    """下载文件"""
    cfg = load_config()
    current_user = _resolve_current_user(authorization, cfg)

    files = get_session_files(session_id, config=cfg, user_id=current_user.id if current_user else None)
    file_info = next((f for f in files if f.id == file_id), None)

    if not file_info:
        raise HTTPException(status_code=404, detail="文件不存在")

    storage_key = getattr(file_info, "storage_key", None) or ""
    file_path = None

    if storage_key and oss_storage_enabled(cfg):
        cache_path = Path(cfg.temp_dir) / "oss_cache" / storage_key
        try:
            file_path = download_file_to_local(storage_key, cache_path, config=cfg)
        except Exception:
            raise HTTPException(status_code=404, detail="文件不存在")
    elif file_info.file_path:
        local_path = Path(file_info.file_path)
        if local_path.exists():
            file_path = local_path
        else:
            raise HTTPException(status_code=404, detail="文件不存在")

    if not file_path:
        raise HTTPException(status_code=404, detail="文件不存在")

    return FileResponse(
        path=str(file_path),
        filename=file_info.file_name,
        media_type="application/octet-stream",
    )


# ============ 通用路径下载（本地测试用）============
# 不依赖 session_id，用于下载 Agent 生成的 _filled.xlsx / _filtered_rows.json 等文件

download_router = APIRouter(prefix="/api/files", tags=["文件下载"])


@download_router.get("/download-by-blob")
async def download_by_blob(blob_name: str, authorization: Optional[str] = Header(default=None)):
    """从阿里云 OSS 按对象键下载文件。"""
    cfg = load_config()
    current_user = _resolve_current_user(authorization, cfg)
    if not oss_storage_enabled(cfg):
        raise HTTPException(status_code=503, detail="OSS 未启用")

    # 允许 workflows / sessions 前缀的 blob
    allowed_prefixes = (cfg.storage.object_key_prefix or "sessions").split(",")
    safe_prefixes = tuple(p.strip() for p in allowed_prefixes) + ("sessions", "workflows")
    if not any(blob_name.startswith(sp) for sp in safe_prefixes):
        raise HTTPException(status_code=403, detail="不允许访问该对象键")

    try:
        from core.storage import download_file_to_local
        import tempfile, os
        ext = os.path.splitext(blob_name)[1] or ""
        fd, tmp_path = tempfile.mkstemp(suffix=ext)
        os.close(fd)
        downloaded = download_file_to_local(blob_name, tmp_path, config=cfg)
        return FileResponse(
            path=str(downloaded),
            filename=downloaded.name,
            media_type="application/octet-stream",
        )
    except Exception:
        raise HTTPException(status_code=404, detail="OSS 对象不存在")


@download_router.get("/download")
async def download_by_path(path: str, authorization: Optional[str] = Header(default=None)):
    """根据本地绝对路径或 OSS 对象键下载（本地测试用）"""
    cfg = load_config()
    current_user = _resolve_current_user(authorization, cfg)

    # 防止路径穿越：只允许 output_dir / temp_dir / library 目录下的文件
    allowed_roots = [Path(cfg.output_dir).resolve(), Path(cfg.temp_dir).resolve(), Path("workspace/library").resolve()]
    file_path = Path(path).resolve()
    if not any(file_path.is_relative_to(root) for root in allowed_roots):
        raise HTTPException(status_code=403, detail="不允许访问该路径")

    if not file_path.exists() and oss_storage_enabled(cfg):
        try:
            cache_path = Path(cfg.temp_dir) / "oss_cache" / path
            file_path = download_file_to_local(path, cache_path, config=cfg)
        except Exception:
            raise HTTPException(status_code=404, detail="文件不存在")
    elif not file_path.exists():
        raise HTTPException(status_code=404, detail="文件不存在")
    return FileResponse(
        path=str(file_path),
        filename=file_path.name,
        media_type="application/octet-stream",
    )


@download_router.get("/preview")
async def preview_file(file_id: str, space_id: str, authorization: Optional[str] = Header(default=None)):
    """预览文档库文件内容（前 N 行文本或表格）。file_id 为文档库中的 UUID。"""
    cfg = load_config()
    current_user = _resolve_current_user(authorization, cfg)

    # 从文档库获取文件信息
    lib = get_desktop_local_library()
    if not lib:
        raise HTTPException(status_code=500, detail="文档库未初始化")

    doc_info = lib.get_doc(space_id, file_id)
    if not doc_info:
        raise HTTPException(status_code=404, detail="文件不存在")

    file_name = doc_info.get("file_name", "")
    # 获取文件实际路径
    storage_key = doc_info.get("storage_key") or ""
    local_path = doc_info.get("local_path") or ""

    file_path = Path(local_path) if local_path and Path(local_path).exists() else None
    if not file_path and storage_key and oss_storage_enabled(cfg):
        try:
            cache_path = Path(cfg.temp_dir) / "oss_cache" / storage_key
            file_path = download_file_to_local(storage_key, cache_path, config=cfg)
        except Exception:
            pass
    if not file_path or not file_path.exists():
        raise HTTPException(status_code=404, detail="文件无法读取")

    ext = file_path.suffix.lower()
    max_lines = 50
    max_table_rows = 20

    try:
        if ext in (".txt", ".md"):
            text = file_path.read_text(encoding="utf-8", errors="replace")
            lines = text.splitlines()[:max_lines]
            return {"type": "text", "content": "\n".join(lines), "total_lines": len(text.splitlines()), "truncated": len(text.splitlines()) > max_lines}

        elif ext in (".xlsx", ".xls"):
            from openpyxl import load_workbook
            wb = load_workbook(str(file_path), read_only=True, data_only=True)
            ws = wb.active
            rows = []
            for idx, row in enumerate(ws.iter_rows(max_row=max_table_rows + 1, values_only=True)):
                rows.append([str(c) if c is not None else "" for c in row])
            wb.close()
            total_rows = ws.max_row or 0
            return {"type": "table", "headers": rows[0] if rows else [], "rows": rows[1:], "total_rows": total_rows, "truncated": total_rows > max_table_rows}

        elif ext == ".docx":
            from docx import Document
            doc = Document(str(file_path))
            paragraphs = []
            for i, para in enumerate(doc.paragraphs[:max_lines]):
                if para.text.strip():
                    style = para.style.name if para.style else ""
                    paragraphs.append({"text": para.text, "style": style})
            return {"type": "text", "content": "\n".join(p["text"] for p in paragraphs), "total_paragraphs": len(doc.paragraphs), "truncated": len(doc.paragraphs) > max_lines}

        elif ext == ".pdf":
            try:
                import pdfplumber
                text_parts = []
                with pdfplumber.open(str(file_path)) as pdf:
                    total_pages = len(pdf.pages)
                    for page in pdf.pages[:3]:
                        page_text = page.extract_text() or ""
                        text_parts.append(page_text)
                    full_text = "\n---\n".join(text_parts)
                return {"type": "text", "content": full_text[:8000], "total_pages": total_pages, "truncated": total_pages > 3}
            except ImportError:
                return {"type": "text", "content": "[PDF 预览需要安装 pdfplumber]", "total_pages": 0, "truncated": False}

        else:
            return {"type": "text", "content": f"[{ext} 格式暂不支持预览]", "total_lines": 0, "truncated": False}

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"预览失败: {e}")


# ============ 临时文件上传（不上传数据库）============

@temp_router.post("/upload", response_model=TempFileResponse)
async def upload_temp_file(
    session_id: str,
    file: UploadFile,
    file_type: str = Form(..., description="文件类型: data 或 template"),
    authorization: Optional[str] = Header(default=None),
):
    """
    上传临时文件（仅保存文件，不存入数据库）
    文件信息返回给前端，前端本地管理
    """
    cfg = load_config()
    _resolve_current_user(authorization, cfg)

    # 验证文件类型
    if file_type not in ("data", "template"):
        raise HTTPException(status_code=400, detail="file_type 必须是 'data' 或 'template'")

    # 保存文件到临时目录
    file_name = file.filename or "unnamed"
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_name = f"{timestamp}_{file_name}"
    file_size = 0

    if oss_storage_enabled(cfg):
        storage_key = upload_stream_to_storage(
            file.file,
            config=cfg,
            blob_name=build_blob_name(session_id, safe_name, prefix=cfg.storage.object_key_prefix),
            content_type=file.content_type,
        )
        if not storage_key:
            raise HTTPException(status_code=502, detail="OSS 上传失败")
        try:
            file.file.seek(0, 2)
            file_size = file.file.tell()
        except Exception:
            file_size = 0
    else:
        file_path = TEMP_UPLOAD_DIR / session_id / safe_name
        file_path.parent.mkdir(parents=True, exist_ok=True)
        with open(file_path, "wb") as buffer:
            while chunk := file.file.read(8192):
                buffer.write(chunk)
                file_size += len(chunk)
        storage_key = str(file_path)

    return TempFileResponse(
        file_name=file_name,
        file_path=storage_key,
        file_type=file_type,
        file_size=file_size,
        storage_key=storage_key,
        created_at=datetime.now().isoformat() + "Z",
    )


@temp_router.delete("/{file_path:path}")
async def delete_temp_file(
    session_id: str,
    file_path: str,
    authorization: Optional[str] = Header(default=None),
):
    """
    删除临时文件
    前端取消上传或删除时调用
    """
    cfg = load_config()
    _resolve_current_user(authorization, cfg)

    path_obj = Path(file_path)
    if not str(path_obj).startswith(str(TEMP_UPLOAD_DIR)):
        raise HTTPException(status_code=400, detail="无效的临时文件路径")

    try:
        if path_obj.exists():
            path_obj.unlink()
    except Exception:
        pass

    # 尝试删除空目录
    try:
        parent = path_obj.parent
        if parent.exists() and not any(parent.iterdir()):
            parent.rmdir()
    except Exception:
        pass

    return {"success": True}
