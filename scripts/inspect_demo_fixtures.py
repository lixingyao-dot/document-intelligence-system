#!/usr/bin/env python3
"""查看课堂演示样例文件结构。"""
from __future__ import annotations

import sys
from pathlib import Path

APP = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(APP / "src"))

import openpyxl
from docx import Document

root = next(p for p in APP.iterdir() if p.is_dir() and (p / "README.txt").exists())
print("ROOT:", root.name)

for sub in sorted(root.iterdir()):
    if not sub.is_dir():
        continue
    print("\n" + "=" * 60)
    print(sub.name)
    req = sub / "用户要求.txt"
    if req.exists():
        print("要求:", req.read_text(encoding="utf-8").strip()[:300])
    for f in sorted(sub.iterdir()):
        if f.suffix.lower() == ".xlsx":
            wb = openpyxl.load_workbook(f, read_only=True, data_only=True)
            print(f"\n  XLSX: {f.name} sheets={wb.sheetnames}")
            for sn in wb.sheetnames[:2]:
                ws = wb[sn]
                rows = []
                for i, row in enumerate(ws.iter_rows(max_row=5, values_only=True), 1):
                    rows.append((i, row))
                print(f"    [{sn}] head:")
                for i, row in rows:
                    print(f"      {i}: {row}")
            wb.close()
        elif f.suffix.lower() == ".docx":
            doc = Document(str(f))
            print(f"\n  DOCX: {f.name} paragraphs={len(doc.paragraphs)} tables={len(doc.tables)}")
            for i, p in enumerate(doc.paragraphs[:8]):
                t = (p.text or "").strip()
                if t:
                    print(f"    p{i}: {t[:120]}")
            for ti, tbl in enumerate(doc.tables[:2]):
                print(f"    table{ti} {len(tbl.rows)}x{len(tbl.columns)}")
                for ri, row in enumerate(tbl.rows[:4]):
                    print(f"      r{ri}: {[c.text.strip() for c in row.cells]}")
